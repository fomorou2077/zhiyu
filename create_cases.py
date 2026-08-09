"""
创建舆情案例库文档
"""
from docx import Document
from datetime import datetime

def create_case_documents():
    # 案例1：地域歧视风波
    doc1 = Document()
    doc1.add_heading('某茶饮品牌广告地域歧视风波', 0)
    doc1.add_paragraph('案例编号：2023-CASE-001')
    doc1.add_paragraph('创建时间：' + datetime.now().strftime("%Y-%m-%d"))
    doc1.add_heading('一、事件概述', level=1)
    doc1.add_paragraph('2023年5月，某茶饮品牌在海南地区的广告宣传中，使用了以下文案："椰子里不喝椰子水，就像海南人不吃椰子鸡"。该文案被网友认为暗示"海南人必须吃椰子鸡"，涉嫌地域歧视。事件在抖音和小红书迅速发酵，24小时内登上微博热搜前三。')

    doc1.add_heading('二、舆情传播路径', level=1)
    doc1.add_paragraph('• 5月18日 10:00 - 首发网友在小红书发布质疑帖')
    doc1.add_paragraph('• 5月18日 14:00 - 话题扩散至微博，引发2000+讨论')
    doc1.add_paragraph('• 5月18日 18:00 - 品牌方发现舆情，启动应急预案')
    doc1.add_paragraph('• 5月19日 00:00 - 品牌发布道歉声明')

    doc1.add_heading('三、企业应对措施', level=1)
    doc1.add_paragraph('1. 6小时内删除所有争议广告物料')
    doc1.add_paragraph('2. 发布官方致歉声明，承认审核不严')
    doc1.add_paragraph('3. 承诺加强内容审核机制')
    doc1.add_paragraph('4. 邀请海南文化专家参与后续宣传设计')

    doc1.add_heading('四、舆论反馈', level=1)
    doc1.add_paragraph('正面：快速响应态度获得部分认可')
    doc1.add_paragraph('负面：仍有部分网友认为道歉不够诚恳')
    doc1.add_paragraph('中性：多数用户关注后续整改措施')

    doc1.add_heading('五、经验教训', level=1)
    doc1.add_paragraph('1. 避免使用地域刻板印象进行营销')
    doc1.add_paragraph('2. 跨界营销需提前进行文化敏感度测试')
    doc1.add_paragraph('3. 建立区域化内容的双重审核机制')
    doc1.add_paragraph('4. 舆情响应需在黄金4小时内完成首次表态')

    doc1.save('static/cases/case1.docx')
    print('案例1创建成功')

    # 案例2：直播带货翻车
    doc2 = Document()
    doc2.add_heading('某直播带货主播夸大功效翻车事件', 0)
    doc2.add_paragraph('案例编号：2024-CASE-002')
    doc2.add_paragraph('创建时间：' + datetime.now().strftime("%Y-%m-%d"))
    doc2.add_heading('一、事件概述', level=1)
    doc2.add_paragraph('2024年1月，某头部带货主播在直播间推销某保健品时，使用了"三天见效"、"无效退款"等宣传用语。职业打假人随即发布测评视频指出产品并无宣传功效，引发消费者集体投诉。')

    doc2.add_heading('二、违规问题点', level=1)
    doc2.add_paragraph('1. 违反《广告法》关于保健食品不得宣传疗效的规定')
    doc2.add_paragraph('2. 使用绝对化用语"三天见效"')
    doc2.add_paragraph('3. 夸大产品功效误导消费者')
    doc2.add_paragraph('4. 未明确标注"广告"字样')

    doc2.add_heading('三、企业应对', level=1)
    doc2.add_paragraph('• 主播团队发布道歉视频')
    doc2.add_paragraph('• 承诺"退一赔三"方案')
    doc2.add_paragraph('• 下架涉事产品')
    doc2.add_paragraph('• 暂停直播一周进行内部整顿')

    doc2.add_heading('四、监管部门处罚', level=1)
    doc2.add_paragraph('市场监管部门立案调查，最终罚款50万元，并要求限期整改直播带货流程。')

    doc2.add_heading('五、经验教训', level=1)
    doc2.add_paragraph('1. 直播话术必须遵守《广告法》')
    doc2.add_paragraph('2. 建立法务事前审核机制')
    doc2.add_paragraph('3. 避免使用绝对化宣传用语')
    doc2.add_paragraph('4. 保健品、药品类需特别谨慎')

    doc2.save('static/cases/case2.docx')
    print('案例2创建成功')

    # 案例3：文案抄袭争议
    doc3 = Document()
    doc3.add_heading('某手机品牌新品海报抄袭争议', 0)
    doc3.add_paragraph('案例编号：2024-CASE-003')
    doc3.add_paragraph('创建时间：' + datetime.now().strftime("%Y-%m-%d"))
    doc3.add_heading('一、事件概述', level=1)
    doc3.add_paragraph('2024年8月，某国产手机品牌发布的新品宣传海报被国外设计师发现与自己的原创作品高度相似，相似度超过90%。设计师在社交媒体发布对比图，指责品牌方"拿来主义"。')

    doc3.add_heading('二、对比分析', level=1)
    doc3.add_paragraph('• 色调搭配：几乎一致')
    doc3.add_paragraph('• 构图方式：采用相同视角')
    doc3.add_paragraph('• 核心元素：手机轮廓，光影效果高度重合')
    doc3.add_paragraph('• 唯一区别：更换了产品型号')

    doc3.add_heading('三、企业应对', level=1)
    doc3.add_paragraph('1. 24小时内删除所有争议海报')
    doc3.add_paragraph('2. 发布官方道歉声明，承认"审核疏忽"')
    doc3.add_paragraph('3. 主动联系原作者，支付版权授权费用')
    doc3.add_paragraph('4. 承诺加强创意素材版权审查')

    doc3.add_heading('四、舆论走向', level=1)
    doc3.add_paragraph('由于品牌方响应迅速、态度诚恳，舆论由最初的负面转为中性偏正面。部分用户对品牌的知错就改表示认可。')

    doc3.add_heading('五、经验教训', level=1)
    doc3.add_paragraph('1. 创意素材需进行严格的版权审查')
    doc3.add_paragraph('2. 建立外部法律顾问团队提前规避风险')
    doc3.add_paragraph('3. 设计外包需签署完整的版权协议')
    doc3.add_paragraph('4. 舆情应对需"快、诚、实"')

    doc3.save('static/cases/case3.docx')
    print('案例3创建成功')

    # 案例4：高管不当言论
    doc4 = Document()
    doc4.add_heading('某车企CEO公开场合发表不当言论', 0)
    doc4.add_paragraph('案例编号：2025-CASE-004')
    doc4.add_paragraph('创建时间：' + datetime.now().strftime("%Y-%m-%d"))
    doc4.add_heading('一、事件概述', level=1)
    doc4.add_paragraph('2025年2月，某新能源车企CEO在新品发布会后接受媒体采访时，被问及如何看待不购买自家产品的消费者时，回答："不买我们品牌汽车的，那都是不懂车的。"该言论被剪辑成短视频疯传，引发网友群嘲。')

    doc4.add_heading('二、争议焦点', level=1)
    doc4.add_paragraph('1. 使用贬低性表述"不懂车"')
    doc4.add_paragraph('2. 暗示只有购买自家产品才算"懂车"')
    doc4.add_paragraph('3. 傲慢态度引发消费者反感')
    doc4.add_paragraph('4. 言论被断章取义广泛传播')

    doc4.add_heading('三、企业应对', level=1)
    doc4.add_paragraph('阶段一（2小时内）：企业官方账号发表声明，称原话被"断章取义"')
    doc4.add_paragraph('阶段二（6小时内）：CEO本人录制道歉视频，强调"尊重每一位消费者的选择"')
    doc4.add_paragraph('阶段三（24小时内）：企业宣布内部开展公关培训，强化高管发言规范')

    doc4.add_heading('四、舆论反馈', level=1)
    doc4.add_paragraph('• 微博话题阅读量突破2亿')
    doc4.add_paragraph('• 部分网友表示"知错能改"，给予理解')
    doc4.add_paragraph('• 部分潜在用户表示"转粉"，不会再考虑该品牌')
    doc4.add_paragraph('• 竞品趁机营销，发布"尊重用户选择"的对比广告')

    doc4.add_heading('五、经验教训', level=1)
    doc4.add_paragraph('1. 高管发言需经公关团队审核')
    doc4.add_paragraph('2. 避免使用绝对化、贬低性表述')
    doc4.add_paragraph('3. 出现争议应第一时间温和回应')
    doc4.add_paragraph('4. 建立高管媒体发言规范手册')

    doc4.save('static/cases/case4.docx')
    print('案例4创建成功')

    print('\n所有案例文档创建完成！')

if __name__ == '__main__':
    create_case_documents()
