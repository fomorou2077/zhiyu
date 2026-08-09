"""
文件处理服务
用于从上传的文件中提取文本内容
"""

import os
import shutil
from pathlib import Path
from typing import List, Optional
from fastapi import UploadFile

from app.utils.logger import logger

# 上传文件存储目录
UPLOAD_DIR = Path(__file__).parent.parent.parent / "uploads" / "plan_analysis"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# 最大文件大小 (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# 支持的文件类型
ALLOWED_EXTENSIONS = {".docx", ".pdf", ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"}


def get_file_extension(filename: str) -> str:
    """获取文件扩展名（小写）"""
    return Path(filename).suffix.lower()


def is_allowed_file(filename: str) -> bool:
    """检查文件类型是否允许"""
    return get_file_extension(filename) in ALLOWED_EXTENSIONS


async def extract_text_from_docx(file_path: str) -> str:
    """
    从 DOCX 文件提取文本

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # 尝试提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        result = "\n".join(text_parts)
        logger.info("从 DOCX 提取文本: {} 字符", len(result))
        return result

    except ImportError:
        logger.warning("python-docx 未安装，无法提取 DOCX 文件")
        return "[DOCX 文件内容无法提取，请确保 python-docx 已安装]"
    except Exception as e:
        logger.error("提取 DOCX 文本失败: {}", e)
        return f"[提取 DOCX 文件失败: {str(e)}]"


async def extract_text_from_pdf(file_path: str) -> str:
    """
    从 PDF 文件提取文本

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    try:
        import pypdf

        text_parts = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"[第{page_num + 1}页]\n{text.strip()}")
                except Exception as e:
                    logger.warning("提取 PDF 第{}页失败: {}", page_num + 1, e)
                    text_parts.append(f"[第{page_num + 1}页 - 内容无法提取]")

        result = "\n\n".join(text_parts)
        logger.info("从 PDF 提取文本: {} 字符, {} 页", len(result), len(reader.pages))
        return result

    except ImportError:
        logger.warning("pypdf 未安装，无法提取 PDF 文件")
        return "[PDF 文件内容无法提取，请确保 pypdf 已安装]"
    except Exception as e:
        logger.error("提取 PDF 文本失败: {}", e)
        return f"[提取 PDF 文件失败: {str(e)}]"


async def extract_text_from_image(file_path: str) -> str:
    """
    从图片文件提取文本（OCR）

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    try:
        import pytesseract
        from PIL import Image

        # 打开图片
        img = Image.open(file_path)

        # 使用 Tesseract OCR 提取文本
        # 可选：先转换为灰度图提高识别率
        if img.mode != "L":
            img = img.convert("L")

        text = pytesseract.image_to_string(img, lang="chi_sim+eng")

        if text and text.strip():
            logger.info("从图片提取 OCR 文本: {} 字符", len(text))
            return text.strip()
        else:
            return "[图片中未识别到文字]"

    except ImportError:
        logger.warning("pytesseract 或 Pillow 未安装，无法提取图片文本")
        return "[图片内容暂无法识别，请安装 pytesseract 和 Pillow]"
    except Exception as e:
        logger.error("提取图片 OCR 文本失败: {}", e)
        return f"[图片 OCR 识别失败: {str(e)}]"


async def process_uploaded_files(files: List[UploadFile], save_dir: Optional[Path] = None) -> str:
    """
    处理上传的文件列表，提取文本并合并

    Args:
        files: 上传的文件列表
        save_dir: 保存文件的目录，None 则使用默认目录

    Returns:
        合并后的文本内容
    """
    if not files:
        return ""

    if save_dir is None:
        save_dir = UPLOAD_DIR

    save_dir.mkdir(parents=True, exist_ok=True)

    text_parts = []

    for idx, file in enumerate(files):
        if not file:
            continue

        filename = file.filename or f"file_{idx}"
        ext = get_file_extension(filename)

        # 检查文件类型
        if not is_allowed_file(filename):
            logger.warning("不支持的文件类型: {}", ext)
            text_parts.append(f"\n[不支持的文件类型: {filename}]\n")
            continue

        # 生成安全文件名
        safe_filename = f"{idx}_{Path(filename).name}"
        file_path = save_dir / safe_filename

        # 保存文件
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            file_size = file_path.stat().st_size
            if file_size > MAX_FILE_SIZE:
                text_parts.append(f"\n[文件过大 ({file_size / 1024 / 1024:.1f}MB > 10MB): {filename}]\n")
                file_path.unlink()  # 删除过大的文件
                continue

            logger.info("处理文件: {}, 大小: {}KB", filename, file_size / 1024)

            # 根据文件类型提取文本
            if ext == ".docx":
                text = await extract_text_from_docx(str(file_path))
            elif ext == ".pdf":
                text = await extract_text_from_pdf(str(file_path))
            elif ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"}:
                text = await extract_text_from_image(str(file_path))
            else:
                text = f"[不支持处理此文件类型: {filename}]"

            text_parts.append(f"\n===== {filename} =====\n{text}\n")

        except Exception as e:
            logger.error("处理文件 {} 失败: {}", filename, e)
            text_parts.append(f"\n[处理文件失败: {filename}]\n")
        finally:
            # 清理临时文件
            if file_path.exists():
                try:
                    file_path.unlink()
                except:
                    pass

    result = "\n".join(text_parts)
    logger.info("文件处理完成，总文本长度: {} 字符", len(result))
    return result


def cleanup_temp_files(directory: Path, max_age_hours: int = 24):
    """
    清理临时文件

    Args:
        directory: 要清理的目录
        max_age_hours: 文件最大保留时间（小时）
    """
    import time

    if not directory.exists():
        return

    current_time = time.time()
    max_age_seconds = max_age_hours * 3600

    for file_path in directory.glob("*"):
        if file_path.is_file():
            file_age = current_time - file_path.stat().st_mtime
            if file_age > max_age_seconds:
                try:
                    file_path.unlink()
                    logger.info("删除过期临时文件: {}", file_path)
                except Exception as e:
                    logger.warning("删除文件失败 {}: {}", file_path, e)
