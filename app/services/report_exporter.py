"""
报告导出服务 - Report Exporter
支持生成 DOCX / PDF 格式的舆情报告
"""
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from app.utils.logger import logger

# 报告模板类型
REPORT_TEMPLATES = {
    "daily": "日报",
    "weekly": "周报",
    "monthly": "月报",
    "incident": "事件专项报告",
    "competitor": "竞品分析报告",
    "plan_risk": "策划风险评估报告",
    "post_event": "事后复盘报告",
}


async def generate_report(
    report_type: str,
    title: str,
    data: dict,
    user_id: int,
    output_format: str = "docx",
) -> dict:
    """
    生成舆情报告
    Args:
        report_type: 报告类型 (daily|weekly|monthly|incident|competitor|plan_risk|post_event)
        title: 报告标题
        data: 报告数据
        user_id: 用户ID
        output_format: 输出格式 (docx|pdf|html)
    Returns:
        {file_path, file_name, format, generated_at}
    """
    template_name = REPORT_TEMPLATES.get(report_type, "舆情分析报告")

    # 创建报告输出目录
    report_dir = Path("reports") / str(user_id)
    report_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c for c in title if c.isalnum() or c in "._- ")[:50]
    file_name = f"{safe_title}_{timestamp}"

    if output_format == "html":
        content = _generate_html_report(template_name, title, data)
        file_path = report_dir / f"{file_name}.html"
        file_path.write_text(content, encoding="utf-8")
        logger.info("HTML报告已生成: {}", file_path)

    elif output_format == "docx":
        content = _generate_docx_content(template_name, title, data)
        file_path = report_dir / f"{file_name}.docx"
        _write_docx(file_path, template_name, title, data)
        logger.info("DOCX报告已生成: {}", file_path)

    elif output_format == "pdf":
        # PDF生成需要额外依赖(如reportlab/weasyprint)，此处生成HTML作为降级方案
        content = _generate_html_report(template_name, title, data)
        file_path = report_dir / f"{file_name}.html"
        file_path.write_text(content, encoding="utf-8")
        logger.warning("PDF生成需要额外依赖（reportlab/weasyprint），已降级为HTML: {}", file_path)
        output_format = "html"

    else:
        raise ValueError(f"不支持的输出格式: {output_format}")

    return {
        "file_path": str(file_path.absolute()),
        "file_name": file_name,
        "format": output_format,
        "generated_at": datetime.now().isoformat(),
        "title": title,
        "template": template_name,
    }


def _generate_html_report(template_name: str, title: str, data: dict) -> str:
    """生成HTML格式的报告"""
    now = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    # 构建情绪分析表格
    sentiment_rows = ""
    if "sentiment" in data:
        sentiment = data["sentiment"]
        sentiment_rows = f"""
        <tr><td>正面</td><td>{sentiment.get('positive', '--')}</td></tr>
        <tr><td>中性</td><td>{sentiment.get('neutral', '--')}</td></tr>
        <tr><td>负面</td><td>{sentiment.get('negative', '--')}</td></tr>"""

    # 关键词
    keywords_html = ""
    if "keywords" in data:
        kw_list = data["keywords"]
        if isinstance(kw_list, list):
            keywords_html = " ".join(f'<span class="tag">{kw}</span>' for kw in kw_list[:20])

    # 提及列表
    mentions_html = ""
    if "mentions" in data:
        for m in data["mentions"][:10]:
            if isinstance(m, dict):
                mentions_html += f"""
                <tr>
                    <td>{m.get('platform', '--')}</td>
                    <td>{m.get('title', '--')}</td>
                    <td>{m.get('sentiment', '--')}</td>
                    <td>{m.get('time', '--')}</td>
                </tr>"""

    # AI分析
    ai_analysis = data.get("ai_analysis", data.get("summary", "暂无AI分析内容"))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{title} - 知舆舆情报告</title>
<style>
  body {{ font-family: "PingFang SC", "Microsoft YaHei", sans-serif; max-width: 800px; margin: 0 auto; padding: 40px 20px; color: #333; }}
  .header {{ text-align: center; border-bottom: 3px solid #00f2ff; padding-bottom: 20px; margin-bottom: 30px; }}
  .header h1 {{ font-size: 28px; margin-bottom: 8px; }}
  .header .meta {{ color: #888; font-size: 14px; }}
  .section {{ margin-bottom: 30px; }}
  .section h2 {{ font-size: 20px; border-left: 4px solid #00f2ff; padding-left: 12px; margin-bottom: 16px; }}
  table {{ width: 100%; border-collapse: collapse; margin-bottom: 16px; }}
  th, td {{ border: 1px solid #e0e0e0; padding: 10px 14px; text-align: left; font-size: 14px; }}
  th {{ background: #f5f5f5; font-weight: 600; }}
  .tag {{ display: inline-block; padding: 3px 10px; background: #e8f4fd; color: #00a8cc; border-radius: 12px; margin: 3px; font-size: 13px; }}
  .analysis-box {{ background: #f8f9fa; border-radius: 8px; padding: 20px; line-height: 1.8; font-size: 15px; }}
  .footer {{ text-align: center; color: #aaa; font-size: 12px; margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; }}
  .stats {{ display: flex; gap: 20px; flex-wrap: wrap; }}
  .stat-card {{ flex: 1; min-width: 120px; background: linear-gradient(135deg, #f0f9ff, #e8f4fd); border-radius: 8px; padding: 16px; text-align: center; }}
  .stat-card .value {{ font-size: 28px; font-weight: 700; color: #00a8cc; }}
  .stat-card .label {{ font-size: 13px; color: #888; margin-top: 4px; }}
</style>
</head>
<body>
<div class="header">
  <h1>{title}</h1>
  <p class="meta">{template_name} · 生成时间：{now} · 知舆ZhiYu舆情监测平台</p>
</div>

<div class="section">
  <h2>📊 数据概览</h2>
  <div class="stats">
    <div class="stat-card"><div class="value">{data.get('total_mentions', '--')}</div><div class="label">总提及量</div></div>
    <div class="stat-card"><div class="value">{data.get('platform_count', '--')}</div><div class="label">覆盖平台</div></div>
    <div class="stat-card"><div class="value">{data.get('risk_level', '--')}</div><div class="label">风险等级</div></div>
  </div>
</div>

<div class="section">
  <h2>🎯 情绪分析</h2>
  <table>{sentiment_rows}</table>
</div>

<div class="section">
  <h2>🔑 热点关键词</h2>
  <p>{keywords_html or '暂无关键词数据'}</p>
</div>

<div class="section">
  <h2>📋 最新提及</h2>
  <table>
    <tr><th>平台</th><th>标题</th><th>情绪</th><th>时间</th></tr>
    {mentions_html or '<tr><td colspan="4">暂无数据</td></tr>'}
  </table>
</div>

<div class="section">
  <h2>🤖 AI智能分析</h2>
  <div class="analysis-box">{ai_analysis}</div>
</div>

<div class="footer">
  <p>本报告由知舆 ZhiYu 舆情监测平台自动生成 · 仅供参考</p>
</div>
</body>
</html>"""
    return html


def _write_docx(file_path: Path, template_name: str, title: str, data: dict):
    """生成DOCX文件（使用python-docx如果可用，否则生成HTML降级）"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        # 标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"{template_name} · 生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')} · 知舆ZhiYu")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 数据概览
        doc.add_heading("数据概览", level=1)
        table = doc.add_table(rows=2, cols=3, style="Table Grid")
        overview_data = [
            ("总提及量", str(data.get("total_mentions", "--"))),
            ("覆盖平台", str(data.get("platform_count", "--"))),
            ("风险等级", str(data.get("risk_level", "--"))),
        ]
        for i, (label, value) in enumerate(overview_data):
            if i < 3:
                table.cell(0, i).text = label
                table.cell(1, i).text = value

        # AI分析
        doc.add_heading("AI智能分析", level=1)
        ai_text = data.get("ai_analysis", data.get("summary", "暂无AI分析内容"))
        doc.add_paragraph(ai_text)

        # 页脚
        doc.add_paragraph("\n---\n本报告由知舆 ZhiYu 舆情监测平台自动生成 · 仅供参考")

        file_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(file_path))
    except ImportError:
        # python-docx不可用时降级为HTML
        html_content = _generate_html_report(template_name, title, data)
        html_path = file_path.with_suffix(".html")
        html_path.write_text(html_content, encoding="utf-8")
        logger.info("python-docx不可用，已生成HTML报告: {}", html_path)


def _generate_docx_content(template_name: str, title: str, data: dict) -> str:
    """生成DOCX内容的文本表示（降级方案）"""
    return _generate_html_report(template_name, title, data)
